"""
PDF / DOCX text extractor for medical report uploads.

Tries pdfplumber first (best for text-layer PDFs).
Falls back to pypdf2 if pdfplumber fails.
DOCX uses python-docx (already in requirements).
"""

import io
from typing import Optional


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF file's bytes.
    Tries pdfplumber first, falls back to pypdf2.
    Raises ValueError if both fail or no text is found.
    """
    text = _try_pdfplumber(file_bytes)
    if text:
        return text

    text = _try_pypdf2(file_bytes)
    if text:
        return text

    raise ValueError(
        "Could not extract text from this PDF. "
        "It may be a scanned image without a text layer. "
        "Please upload a text-based PDF or a DOCX file."
    )


def _try_pdfplumber(file_bytes: bytes) -> Optional[str]:
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = []
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    pages.append(page_text.strip())
            text = "\n\n".join(pages).strip()
            return text if text else None
    except ImportError:
        return None
    except Exception:
        return None


def _try_pypdf2(file_bytes: bytes) -> Optional[str]:
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        pages = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                pages.append(page_text.strip())
        text = "\n\n".join(pages).strip()
        return text if text else None
    except ImportError:
        pass
    except Exception:
        pass

    # also try the older PyPDF2 package name
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        pages = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                pages.append(page_text.strip())
        text = "\n\n".join(pages).strip()
        return text if text else None
    except Exception:
        return None


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX file's bytes using python-docx.
    Raises ValueError if extraction fails.
    """
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = []

        # Paragraphs
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)

        # Table cells (same pattern as docx_ingest.py)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t:
                        parts.append(t)

        text = "\n".join(parts).strip()
        if not text:
            raise ValueError("No text content found in this DOCX file.")
        return text

    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"Failed to read DOCX file: {e}")


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Route to the correct extractor based on file extension.
    Returns extracted text string.
    Raises ValueError for unsupported types or extraction failures.
    """
    lower = filename.lower()

    if lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)

    if lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)

    raise ValueError(
        f"Unsupported file type for '{filename}'. "
        "Please upload a PDF (.pdf) or Word document (.docx)."
    )
