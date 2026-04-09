"""
Extract text from .docx files and ingest into RAG for psychologist chatbot.

Use this to load psychology / assessment documents (e.g. Intelligence Test,
Personality Profiler, Emotional Intelligence Scale, Peer Relationship Test)
into the knowledge base so Uma can use them when supporting users.
"""

from pathlib import Path
from typing import Optional

from docx import Document


def extract_text_from_docx(path: Path) -> str:
    """Extract all text from a .docx file (paragraphs + table cells)."""
    doc = Document(str(path))
    parts = []

    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = (cell.text or "").strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                parts.append(" | ".join(row_text))

    return "\n\n".join(parts) if parts else ""


def ingest_docx_folder(
    folder: str | Path,
    rag_store=None,
    *,
    pattern: str = "*.docx",
) -> tuple[int, list[str], list[str]]:
    """
    Load all .docx files from a folder into the RAG store.

    Returns:
        (files_processed, chunk_ids, errors)
    """
    if rag_store is None:
        from dotenv import load_dotenv
        load_dotenv()
        from rag import get_rag_store
        rag_store = get_rag_store()

    folder = Path(folder)
    if not folder.is_dir():
        return 0, [], [f"Not a directory: {folder}"]

    files = list(folder.glob(pattern))
    all_ids = []
    errors = []

    for path in sorted(files):
        try:
            text = extract_text_from_docx(path)
            if not text.strip():
                errors.append(f"{path.name}: no text extracted")
                continue
            meta = {"source": path.name, "type": "psychology_doc"}
            ids = rag_store.add_documents(
                [text],
                metadata_per_doc=[meta],
                auto_chunk=True,
            )
            all_ids.extend(ids)
        except Exception as e:
            errors.append(f"{path.name}: {e}")

    return len(files), all_ids, errors


def main():
    """CLI: python -m docx_ingest [folder_path]"""
    import sys
    from dotenv import load_dotenv
    load_dotenv()

    if not __import__("os").environ.get("OPENAI_API_KEY"):
        print("ERROR: Set OPENAI_API_KEY in .env (needed for embeddings)")
        sys.exit(1)

    folder = sys.argv[1] if len(sys.argv) > 1 else None
    if not folder:
        default = Path(__file__).parent / "data" / "psychology_docs"
        default.mkdir(parents=True, exist_ok=True)
        print(f"Usage: python -m docx_ingest <folder_with_docx>")
        print(f"Example: python -m docx_ingest \"C:\\...\\transfers\\2026-11\"")
        print(f"Or copy .docx files into: {default}")
        print(f"Then run: python -m docx_ingest \"{default}\"")
        sys.exit(0)

    n, ids, errs = ingest_docx_folder(folder)
    print(f"Processed {n} file(s), added {len(ids)} chunk(s) to RAG.")
    if errs:
        for e in errs:
            print(f"  - {e}")
    if ids:
        print("Psychologist docs are now in the knowledge base. Restart the API if it was running.")


if __name__ == "__main__":
    main()
