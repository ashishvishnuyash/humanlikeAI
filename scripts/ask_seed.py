"""Seed-and-test helper for /ask: uploads two documents (a synthesized health
checkup DOCX with bloodwork + the on-disk resume DOCX) so /ask has both medical
and personal-document content to retrieve from.

Run BEFORE `python -m scripts.ask_test`:
    python -m scripts.ask_seed --reset    # clears Pinecone + tears down old seed
    python -m scripts.ask_seed            # creates a fresh seed + uploads docs
"""
from __future__ import annotations

import io
import sys
import time
import uuid as _uuid
from pathlib import Path

from fastapi.testclient import TestClient

from main import app
from scripts import api_test_seed


def _build_health_docx() -> bytes:
    from docx import Document
    d = Document()
    d.add_heading("Quarterly Health Check-up Report", 1)
    d.add_paragraph("Patient: Smoke Test Employee")
    d.add_paragraph("Date of report: 2026-05-03")
    d.add_paragraph("Issuing facility: Diltak Wellness Clinic")
    d.add_heading("Vital signs", 2)
    d.add_paragraph("Blood pressure: 118/76 mmHg (normal)")
    d.add_paragraph("Resting heart rate: 64 bpm (normal)")
    d.add_paragraph("BMI: 23.4 (normal range)")
    d.add_paragraph("Body temperature: 36.8 C")
    d.add_heading("Bloodwork", 2)
    d.add_paragraph(
        "Hemoglobin: 14.2 g/dL (reference 13.0-17.0). White blood cells 6.8 x 10^9/L. "
        "Platelets 245 x 10^9/L."
    )
    d.add_paragraph(
        "Lipid panel: Total cholesterol 178 mg/dL, LDL 92 mg/dL, HDL 56 mg/dL, "
        "triglycerides 105 mg/dL."
    )
    d.add_paragraph("Fasting glucose: 88 mg/dL (normal). HbA1c: 5.2%.")
    d.add_heading("Recommendations", 2)
    d.add_paragraph(
        "Continue current diet and exercise routine. Aim for 150 minutes of "
        "moderate aerobic activity per week. Increase water intake to 2.5 L daily. "
        "Schedule next routine check-up in 12 months."
    )
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _wait_for_chunks(doc_id: str, max_wait_s: float = 45.0) -> int:
    from db.models.physical_health import MedicalDocument
    from db.session import get_session_factory
    SessionLocal = get_session_factory()
    deadline = time.time() + max_wait_s
    while time.time() < deadline:
        with SessionLocal() as db:
            row = db.query(MedicalDocument).filter(
                MedicalDocument.id == _uuid.UUID(doc_id)
            ).one_or_none()
            if row and row.rag_chunk_ids:
                return len(row.rag_chunk_ids)
        time.sleep(2.0)
    return 0


def _wipe_pinecone() -> None:
    import os
    from dotenv import load_dotenv
    load_dotenv(".env")
    from pinecone import Pinecone
    api_key = os.environ["PINECONE_API_KEY"]
    host = os.environ.get(
        "PINECONE_HOST",
        "https://diltak-w89a9hg.svc.aped-4627-b74a.pinecone.io",
    )
    pc = Pinecone(api_key=api_key)
    idx = pc.Index(host=host)
    try:
        idx.delete(delete_all=True)
        print(f"Pinecone wiped — total_vectors: {idx.describe_index_stats().get('total_vector_count')}")
    except Exception as e:
        # Empty index throws — that's fine
        if "Namespace not found" in str(e) or "404" in str(e):
            print("Pinecone already empty.")
        else:
            raise


def main() -> int:
    if "--reset" in sys.argv:
        print("Resetting environment...")
        try:
            td = api_test_seed.teardown()
            print(f"  teardown: {td}")
        except Exception as e:
            print(f"  teardown skipped: {e}")
        _wipe_pinecone()
        if "--only-reset" in sys.argv or len(sys.argv) > 2:
            return 0

    print("\nCreating fresh seed user…")
    seeds = api_test_seed.setup(suffix=None)
    emp = seeds["employee"]
    token = emp["access_token"]
    print(f"  employee: {emp['email']}  uid={emp['uid']}")

    client = TestClient(app)

    # 1. Health DOCX (bloodwork)
    print("\nUploading synthesized health DOCX (bloodwork + cholesterol + BP)...")
    health = _build_health_docx()
    r = client.post(
        "/api/physical-health/medical/upload?report_type=blood_test",
        headers={"Authorization": f"Bearer {token}"},
        files={"file": ("checkup_report.docx", health,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    if r.status_code != 202:
        print(f"  FAIL: HTTP {r.status_code} {r.text[:200]}")
        return 1
    health_doc_id = r.json()["doc_id"]
    print(f"  doc_id={health_doc_id}")

    # 2. Resume DOCX
    print("\nUploading Vishal_Singh_Resume.docx...")
    resume_path = Path(__file__).parent / "Vishal_Singh_Resume.docx"
    r = client.post(
        "/api/physical-health/medical/upload?report_type=other",
        headers={"Authorization": f"Bearer {token}"},
        files={"file": (resume_path.name, resume_path.read_bytes(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    if r.status_code != 202:
        print(f"  FAIL: HTTP {r.status_code} {r.text[:200]}")
        return 1
    resume_doc_id = r.json()["doc_id"]
    print(f"  doc_id={resume_doc_id}")

    # 3. Wait for both to land in Pinecone
    print("\nWaiting for RAG ingestion to complete...")
    h_chunks = _wait_for_chunks(health_doc_id)
    r_chunks = _wait_for_chunks(resume_doc_id)
    print(f"  health  : {h_chunks} chunk(s)")
    print(f"  resume  : {r_chunks} chunk(s)")

    if h_chunks == 0 or r_chunks == 0:
        print("\nWARNING: One or both docs failed to ingest into Pinecone.")
        return 1

    print(f"\nReady for /ask testing.")
    print(f"  Run: python -m scripts.ask_test")
    return 0


if __name__ == "__main__":
    sys.exit(main())
