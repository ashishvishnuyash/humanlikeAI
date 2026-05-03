"""Focused smoke test for /api/physical-health endpoints.

Targets the recent changes on origin/main (commits 3536223..431f03e):
  - report generation now requires only 1 check-in (was 3)
  - upload size limit raised to 100 MB
  - rag_chunk_ids column added; delete_medical_document cleans Pinecone

Flow
----
1. seed.setup() → get an employee account + JWT
2. POST /check-in (×2) — populate enough data
3. GET  /check-ins   — verify history
4. GET  /score       — verify composite score
5. GET  /trends?period=7d
6. POST /reports/generate (days=7) — verify it works with <3 check-ins (NEW)
7. GET  /reports     — confirm the just-generated report shows up
8. GET  /medical     — list docs (sanity, expect empty)
9. seed.teardown()

Run:  python -m scripts.physical_health_smoke
"""
from __future__ import annotations

import json
import sys
import time
import traceback
import uuid as _uuid
from typing import Any

from fastapi.testclient import TestClient

from main import app
from scripts import api_test_seed


def _hdr(token: str) -> dict[str, str]:
    return {"Authorization": f"Bearer {token}"}


def _short(payload: Any, n: int = 400) -> Any:
    text = json.dumps(payload, default=str) if not isinstance(payload, str) else payload
    return text if len(text) <= n else text[: n - 3] + "..."


def _check(label: str, resp, expected: set[int]) -> dict:
    ok = resp.status_code in expected
    try:
        body = resp.json()
    except Exception:
        body = resp.text
    print(f"  [{'PASS' if ok else 'FAIL'}] {label} -> HTTP {resp.status_code}")
    if not ok:
        print(f"         body: {_short(body, 600)}")
    return {"label": label, "status": resp.status_code, "ok": ok, "body": body}


def main() -> int:
    keep = "--keep" in sys.argv  # skip DELETE so the Pinecone vector stays for inspection

    print("=" * 70)
    print(
        "PHYSICAL HEALTH SMOKE — testing recent changes on origin/main"
        + (" [--keep mode]" if keep else "")
    )
    print("=" * 70)

    seeds = None
    results: list[dict] = []
    try:
        seeds = api_test_seed.setup(suffix=None)
        emp = seeds["employee"]
        token = emp["access_token"]
        print(f"\nemployee uid: {emp['uid']}\nemail:        {emp['email']}\n")

        client = TestClient(app)

        # 1. POST check-in #1
        body1 = {
            "energy_level": 7, "sleep_quality": 8, "sleep_hours": 7.5,
            "exercise_done": True, "exercise_minutes": 30, "exercise_type": "walk",
            "nutrition_quality": 7, "pain_level": 9, "hydration": 7,
            "notes": "Smoke test check-in #1",
        }
        r = client.post("/api/physical-health/check-in", json=body1, headers=_hdr(token))
        results.append(_check("POST /check-in #1", r, {201}))

        # 2. POST check-in #2
        body2 = {**body1, "energy_level": 6, "notes": "Smoke test check-in #2"}
        r = client.post("/api/physical-health/check-in", json=body2, headers=_hdr(token))
        results.append(_check("POST /check-in #2", r, {201}))

        # 3. GET history
        r = client.get("/api/physical-health/check-ins?days=7", headers=_hdr(token))
        results.append(_check("GET  /check-ins?days=7", r, {200}))
        if r.status_code == 200:
            data = r.json()
            print(f"         total check-ins: {data.get('total')}, returned: {len(data.get('checkins', []))}")

        # 4. GET score
        r = client.get("/api/physical-health/score", headers=_hdr(token))
        results.append(_check("GET  /score", r, {200}))
        if r.status_code == 200:
            d = r.json()
            print(f"         score={d.get('score')} level={d.get('level')} streak={d.get('streak_days')}")

        # 5. GET trends
        r = client.get("/api/physical-health/trends?period=7d", headers=_hdr(token))
        results.append(_check("GET  /trends?period=7d", r, {200}))

        # 6. POST report generate (KEY: previously required ≥3 check-ins, now ≥1)
        r = client.post(
            "/api/physical-health/reports/generate",
            json={"report_type": "on_demand", "days": 7},
            headers=_hdr(token),
        )
        results.append(_check(
            "POST /reports/generate (NEW: <3 check-ins now allowed)",
            r, {201},
        ))
        report_id = None
        if r.status_code == 201:
            d = r.json()
            report_id = d.get("report_id")
            print(f"         report_id={report_id} score={d.get('overall_score')} trend={d.get('trend')}")

        # 7. GET reports list
        r = client.get("/api/physical-health/reports", headers=_hdr(token))
        results.append(_check("GET  /reports", r, {200}))
        if r.status_code == 200:
            d = r.json()
            print(f"         reports total={d.get('total')}")

        # 8. GET single report (if we got one)
        if report_id:
            r = client.get(f"/api/physical-health/reports/{report_id}", headers=_hdr(token))
            results.append(_check(f"GET  /reports/{report_id[:8]}…", r, {200}))

        # 9. GET medical documents (sanity — expect empty)
        r = client.get("/api/physical-health/medical", headers=_hdr(token))
        results.append(_check("GET  /medical (pre-upload, expect empty)", r, {200}))

        # 10. POST /medical/upload — synthesize a DOCX with real text content
        # (PDF parsers cannot extract from image-only PDFs; DOCX has a real text layer)
        import io as _io
        from docx import Document as _Doc
        d = _Doc()
        d.add_heading("Quarterly Health Check-up Report", level=1)
        d.add_paragraph("Patient: Smoke Test Employee")
        d.add_paragraph("Date of report: 2026-05-03")
        d.add_paragraph("Issuing facility: Diltak Wellness Clinic")

        d.add_heading("Vital signs", level=2)
        d.add_paragraph("Blood pressure: 118/76 mmHg (normal)")
        d.add_paragraph("Resting heart rate: 64 bpm (normal)")
        d.add_paragraph("BMI: 23.4 (normal range)")
        d.add_paragraph("Body temperature: 36.8 C")

        d.add_heading("Bloodwork", level=2)
        d.add_paragraph(
            "Complete blood count was within normal limits. Hemoglobin 14.2 g/dL "
            "(reference 13.0-17.0). White blood cell count 6.8 x 10^9/L. "
            "Platelet count 245 x 10^9/L."
        )
        d.add_paragraph(
            "Lipid panel: Total cholesterol 178 mg/dL, LDL 92 mg/dL, HDL 56 mg/dL, "
            "triglycerides 105 mg/dL. All values within target range."
        )
        d.add_paragraph(
            "Fasting glucose: 88 mg/dL (normal). HbA1c: 5.2% (non-diabetic range)."
        )

        d.add_heading("Recommendations", level=2)
        d.add_paragraph(
            "Continue current diet and exercise routine. Recommended at least "
            "150 minutes of moderate aerobic activity per week."
        )
        d.add_paragraph(
            "Increase daily water intake to approximately 2.5 liters."
        )
        d.add_paragraph(
            "Schedule next routine check-up in 12 months. No follow-up consultations "
            "needed at this time. Patient is in excellent overall health."
        )

        d.add_heading("Notes", level=2)
        d.add_paragraph(
            "Patient reports occasional mild lower back tension after long periods "
            "at the desk. Suggested ergonomic workstation review and incorporation "
            "of brief stretching every 90 minutes."
        )

        buf = _io.BytesIO()
        d.save(buf)
        docx_bytes = buf.getvalue()
        print(f"\n  using synthesized DOCX with real text content ({len(docx_bytes):,} bytes)")

        upload_resp = client.post(
            "/api/physical-health/medical/upload?report_type=general_checkup",
            headers=_hdr(token),
            files={
                "file": (
                    "checkup_report.docx",
                    docx_bytes,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            },
        )
        results.append(_check("POST /medical/upload (DOCX, general_checkup)", upload_resp, {202}))
        doc_id = None
        if upload_resp.status_code == 202:
            doc_id = upload_resp.json().get("doc_id")
            print(f"         doc_id={doc_id} status={upload_resp.json().get('status')}")

        # 11. GET /medical/{doc_id}/status (immediately, before background analysis finishes)
        if doc_id:
            r = client.get(f"/api/physical-health/medical/{doc_id}/status", headers=_hdr(token))
            results.append(_check("GET  /medical/{id}/status", r, {200}))

        # 12. GET /medical/{doc_id} detail
        if doc_id:
            r = client.get(f"/api/physical-health/medical/{doc_id}", headers=_hdr(token))
            results.append(_check("GET  /medical/{id} detail", r, {200}))

        # 13. Wait for background analysis (RAG ingestion -> chunk_ids persisted)
        # Then verify the DB row got rag_chunk_ids populated.
        if doc_id:
            from db.session import get_session_factory
            from db.models.physical_health import MedicalDocument
            SessionLocal = get_session_factory()
            chunk_ids: list = []
            for attempt in range(12):  # up to ~30s
                time.sleep(2.5)
                with SessionLocal() as _db:
                    row = (
                        _db.query(MedicalDocument)
                        .filter(MedicalDocument.id == _uuid.UUID(doc_id))
                        .one_or_none()
                    )
                    if row is not None and row.rag_chunk_ids:
                        chunk_ids = list(row.rag_chunk_ids)
                        break
            ok = bool(chunk_ids)
            print(
                f"  [{'PASS' if ok else 'FAIL'}] rag_chunk_ids persisted after upload "
                f"-> {len(chunk_ids)} chunk(s)"
            )
            results.append({"label": "rag_chunk_ids persisted", "status": 0, "ok": ok, "body": chunk_ids[:3]})

        # 14. DELETE /medical/{doc_id} — should clean Pinecone vectors using rag_chunk_ids
        #     Skipped when --keep is passed so the vector remains for inspection.
        if doc_id and not keep:
            r = client.delete(f"/api/physical-health/medical/{doc_id}", headers=_hdr(token))
            results.append(_check("DELETE /medical/{id}", r, {200}))
            if r.status_code == 200:
                msg = r.json().get("message", "")
                print(f"         message: {msg[:200]}")

            # 15. GET medical (post-delete, expect doc gone)
            r = client.get("/api/physical-health/medical", headers=_hdr(token))
            results.append(_check("GET  /medical (post-delete)", r, {200}))
            if r.status_code == 200:
                d2 = r.json()
                still_there = any(x.get("doc_id") == doc_id for x in d2.get("documents", []))
                ok = not still_there
                print(f"         doc gone after delete: {ok}")
                results.append({
                    "label": "document removed after DELETE",
                    "status": 0, "ok": ok, "body": None,
                })
        elif keep and doc_id:
            print(f"\n  [KEEP] Skipping DELETE — Pinecone vector + Postgres row preserved.")
            print(f"         doc_id    : {doc_id}")
            print(f"         chunk_ids : {chunk_ids}")
            print(f"         To clean up later, run:")
            print(f"           curl -X DELETE -H 'Authorization: Bearer <token>' \\")
            print(f"                http://localhost:8000/api/physical-health/medical/{doc_id}")

        print("\n" + "-" * 70)
        passed = sum(1 for x in results if x["ok"])
        failed = len(results) - passed
        print(f"SUMMARY: {passed}/{len(results)} passed, {failed} failed")
        print("-" * 70)
        return 0 if failed == 0 else 1

    except Exception as e:
        print(f"\nFATAL: {e}")
        traceback.print_exc()
        return 2
    finally:
        if seeds is not None:
            if keep:
                # Preserve seeded user + medical_document so the orphaned Pinecone
                # vector can be cleaned later via the DELETE endpoint.
                emp = seeds["employee"]
                print("\n[KEEP] Skipping teardown — seed user preserved.")
                print(f"        employee email: {emp['email']}")
                print(f"        access_token  : {emp['access_token']}")
                print(f"        Manual teardown: python -m scripts.api_test_seed --teardown")
            else:
                try:
                    td = api_test_seed.teardown()
                    print(f"\nteardown: {td}")
                except Exception as e:
                    print(f"teardown failed: {e}")


if __name__ == "__main__":
    sys.exit(main())
